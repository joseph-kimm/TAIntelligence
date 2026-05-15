import asyncio
import io
import logging
import time

import asyncpg
import tiktoken
from docx import Document as DocxDocument
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document as LlamaDocument
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from pypdf import PdfReader
from core.config import settings
from db.chunks import bulk_insert_parent_child_chunks
from db.documents import mark_document_ingestion_failed, update_document_after_ingestion

logger = logging.getLogger(__name__)


def create_embed_model() -> HuggingFaceEmbedding:
    """Load the BGE embedding model. Called once at server startup."""
    return HuggingFaceEmbedding(
        model_name="BAAI/bge-small-en-v1.5",
        token=settings.hf_token
    )


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """Pull plain text out of PDF, DOCX, or TXT bytes."""
    if mime_type == "application/pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = file_bytes.decode("utf-8", errors="replace")
    return text.replace("\x00", "")

def build_chunks(text: str) -> tuple[list[str], list[str], list[int]]:
    """Build non-overlapping parent chunks and overlapping child chunks.

    Parents are 1400-token non-overlapping windows split from the original text.
    Children are 350-token/60-overlap sub-chunks scoped within each parent.
    Returns (parent_texts, child_texts, child_parent_indices).
    """
    parent_splitter = SentenceSplitter(chunk_size=1400, chunk_overlap=0)
    child_splitter = SentenceSplitter(chunk_size=350, chunk_overlap=60)

    parent_nodes = parent_splitter.get_nodes_from_documents([LlamaDocument(text=text)])
    parent_texts = [n.get_content() for n in parent_nodes]

    child_texts: list[str] = []
    child_parent_indices: list[int] = []
    for i, parent_text in enumerate(parent_texts):
        children = child_splitter.get_nodes_from_documents([LlamaDocument(text=parent_text)])
        child_texts.extend(n.get_content() for n in children)
        child_parent_indices.extend([i] * len(children))

    return parent_texts, child_texts, child_parent_indices


def embed_chunks(embed_model: HuggingFaceEmbedding, texts: list[str]) -> list[list[float]]:
    """Return a 384-dim embedding vector for each text string."""
    return embed_model.get_text_embedding_batch(texts, show_progress=False)


async def ingest_document(
    pool: asyncpg.Pool,
    embed_model: HuggingFaceEmbedding,
    document_id: str,
    file_bytes: bytes,
    mime_type: str,
    title: str = "",
) -> None:
    """
    Full ingestion pipeline: extract → chunk → embed → store.
    Runs as a FastAPI BackgroundTask after the upload response is sent.
    """
    label = f'"{title}"' if title else document_id
    t_start = time.perf_counter()

    try:
        logger.info("[%s] Ingestion started (%.1f KB, %s)", label, len(file_bytes) / 1024, mime_type)

        t_extract = time.perf_counter()
        text = await asyncio.to_thread(extract_text, file_bytes, mime_type)
        if not text.strip():
            logger.warning("[%s] No extractable text — skipping", label)
            return

        char_count = len(text)
        word_count = len(text.split())
        token_count = len(tiktoken.get_encoding("cl100k_base").encode(text))
        logger.info("[%s] Extracted text in %.1fs — %d chars, ~%d words, %d tokens", label, time.perf_counter() - t_extract, char_count, word_count, token_count)

        t_chunk = time.perf_counter()
        parent_texts, child_texts, child_parent_indices = await asyncio.to_thread(build_chunks, text)
        if not child_texts:
            logger.warning("[%s] No chunks produced — skipping", label)
            return

        child_sizes = [len(t.split()) for t in child_texts]
        logger.info(
            "[%s] Chunked in %.1fs — %d parents, %d children, min=%d words, avg=%d words, max=%d words",
            label, time.perf_counter() - t_chunk, len(parent_texts), len(child_texts),
            min(child_sizes), sum(child_sizes) // len(child_sizes), max(child_sizes),
        )

        enc = tiktoken.get_encoding("cl100k_base")
        child_token_counts = [len(enc.encode(t)) for t in child_texts]

        logger.info("[%s] Embedding %d child chunks (%d parents)…", label, len(child_texts), len(parent_texts))
        t_embed = time.perf_counter()
        embeddings: list[list[float]] = await asyncio.to_thread(embed_chunks, embed_model, child_texts)
        logger.info("[%s] Embeddings done in %.1fs (dim=%d)", label, time.perf_counter() - t_embed, len(embeddings[0]) if embeddings else 0)

        t_db = time.perf_counter()
        count = await bulk_insert_parent_child_chunks(
            pool, document_id, parent_texts, child_texts, child_token_counts, child_parent_indices, embeddings
        )
        logger.info("[%s] DB write done in %.1fs — %d child chunks, %d parents stored", label, time.perf_counter() - t_db, count, len(parent_texts))

        await update_document_after_ingestion(pool, document_id, token_count, text)

        logger.info("[%s] Ingestion complete in %.1fs total", label, time.perf_counter() - t_start)

    except Exception:
        logger.exception("[%s] Ingestion failed after %.1fs", label, time.perf_counter() - t_start)
        try:
            await mark_document_ingestion_failed(pool, document_id)
        except Exception:
            logger.exception("[%s] Failed to mark document as failed", label)
